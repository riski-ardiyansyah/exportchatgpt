import streamlit as st
import requests
from bs4 import BeautifulSoup
import pandas as pd
import json
from io import BytesIO, StringIO
from fpdf import FPDF
from docx import Document

st.set_page_config(page_title="Export ChatGPT Share", layout="wide")
st.title("🔗 Export ChatGPT Share Link")

url = st.text_input("Masukkan Link Share ChatGPT (format: https://chat.openai.com/share/...)")

export_format = st.selectbox("Pilih Format Ekspor", ["TXT", "CSV", "JSON", "PDF", "DOCX"])

def fetch_chat_from_share_link(url):
    try:
        headers = {"User-Agent": "Mozilla/5.0"}
        response = requests.get(url, headers=headers)
        soup = BeautifulSoup(response.text, "html.parser")
        elements = soup.select('[data-message-author-role]')
        dialog = []

        for el in elements:
            role = el['data-message-author-role'].capitalize()
            message = el.get_text(separator="\n").strip()
            dialog.append({"speaker": role, "message": message})

        return dialog
    except Exception as e:
        st.error(f"Gagal mengambil data dari link: {e}")
        return []

if st.button("Ekspor"):
    if not url.strip():
        st.warning("Harap masukkan link share dari ChatGPT.")
    elif "chat.openai.com/share/" not in url:
        st.error("Link tidak valid. Gunakan link share dari ChatGPT.")
    else:
        chat_data = fetch_chat_from_share_link(url)

        if chat_data:
            if export_format == "TXT":
                output = "\n\n".join([f"{d['speaker']}: {d['message']}" for d in chat_data])
                st.download_button("⬇️ Download .txt", data=output, file_name="chat_export.txt", mime="text/plain")

            elif export_format == "CSV":
                df = pd.DataFrame(chat_data)
                csv = df.to_csv(index=False)
                st.download_button("⬇️ Download .csv", data=csv, file_name="chat_export.csv", mime="text/csv")

            elif export_format == "JSON":
                json_data = json.dumps(chat_data, indent=2, ensure_ascii=False)
                st.download_button("⬇️ Download .json", data=json_data, file_name="chat_export.json", mime="application/json")

            elif export_format == "PDF":
                pdf = FPDF()
                pdf.add_page()
                pdf.set_auto_page_break(auto=True, margin=15)
                pdf.set_font("Arial", size=12)

                for item in chat_data:
                    text = f"{item['speaker']}: {item['message']}\n\n"
                    pdf.multi_cell(0, 10, text)

                pdf_buffer = BytesIO()
                pdf.output(pdf_buffer)
                pdf_buffer.seek(0)
                st.download_button("⬇️ Download .pdf", data=pdf_buffer, file_name="chat_export.pdf", mime="application/pdf")

            elif export_format == "DOCX":
                doc = Document()
                doc.add_heading("ChatGPT Conversation Export", 0)

                for item in chat_data:
                    doc.add_paragraph(f"{item['speaker']}: {item['message']}")

                docx_buffer = BytesIO()
                doc.save(docx_buffer)
                docx_buffer.seek(0)
                st.download_button("⬇️ Download .docx", data=docx_buffer, file_name="chat_export.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        else:
            st.error("Tidak ditemukan obrolan di link tersebut.")
